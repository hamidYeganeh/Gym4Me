from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("/Users/mahdi/Documents/projects/Gym4Me/output/Hamid_Yeganeh_Frontend_Resume_Professional.docx")

# Base: decision_memo (Arial) with named ATS resume overrides.
# Overrides: single-column layout, no tables, no headers/footers, 0.68-inch
# margins, compact type scale, a navy/teal editorial hierarchy, and a light
# impact callout. All identity and contact data remain in the document body.
TOKENS = {
    "page_width": 8.5,
    "page_height": 11.0,
    "margin": 0.68,
    "font": "Arial",
    "body_size": 9.45,
    "body_after": 1.8,
    "body_line": 1.035,
    "title_size": 24.0,
    "role_size": 12.0,
    "section_size": 10.7,
    "entry_size": 10.1,
    "navy": "16324F",
    "teal": "1F6F78",
    "ink": "182433",
    "muted": "536574",
    "rule": "9DB8C0",
    "section_fill": "EEF5F6",
    "impact_fill": "F4F8F8",
    "link": "1C5D75",
}


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = TOKENS["font"]
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), TOKENS["font"])
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), TOKENS["font"])
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_keep_with_next(paragraph, keep=True):
    p_pr = paragraph._p.get_or_add_pPr()
    node = p_pr.find(qn("w:keepNext"))
    if keep and node is None:
        p_pr.append(OxmlElement("w:keepNext"))
    elif not keep and node is not None:
        p_pr.remove(node)


def set_keep_together(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:keepLines")) is None:
        p_pr.append(OxmlElement("w:keepLines"))


def set_cant_split(paragraph):
    set_keep_together(paragraph)


def add_bottom_border(paragraph, color=None, size="6", space="3"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color or TOKENS["rule"])
    p_bdr.append(bottom)


def add_left_border(paragraph, color=None, size="14", space="7"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), space)
    left.set(qn("w:color"), color or TOKENS["teal"])
    p_bdr.append(left)


def add_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)


def add_hyperlink(paragraph, text, url, color=None, bold=False):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), TOKENS["font"])
    r_fonts.set(qn("w:hAnsi"), TOKENS["font"])
    r_pr.append(r_fonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color or TOKENS["link"])
    r_pr.append(c)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(int(TOKENS["body_size"] * 2)))
    r_pr.append(size)
    if bold:
        r_pr.append(OxmlElement("w:b"))
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, pieces, size=None, color=None):
    for piece in pieces:
        if "url" in piece:
            add_hyperlink(
                paragraph,
                piece["text"],
                piece["url"],
                color=piece.get("color", color),
                bold=piece.get("bold", False),
            )
            continue
        run = paragraph.add_run(piece["text"])
        set_run_font(
            run,
            size=piece.get("size", size or TOKENS["body_size"]),
            color=piece.get("color", color or TOKENS["ink"]),
            bold=piece.get("bold"),
            italic=piece.get("italic"),
        )


def add_section_heading(doc, text, page_break_before=False):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(7.0)
    p.paragraph_format.space_after = Pt(4.2)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.page_break_before = page_break_before
    set_keep_with_next(p)
    run = p.add_run(text.upper())
    set_run_font(run, size=TOKENS["section_size"], color=TOKENS["navy"], bold=True)
    add_left_border(p)
    add_paragraph_shading(p, TOKENS["section_fill"])
    return p


def add_body(doc, text="", before=0, after=None, bold_prefix=None, italic=False):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(TOKENS["body_after"] if after is None else after)
    p.paragraph_format.line_spacing = TOKENS["body_line"]
    if bold_prefix and text.startswith(bold_prefix):
        lead = p.add_run(bold_prefix)
        set_run_font(lead, bold=True, size=TOKENS["body_size"], color=TOKENS["ink"])
        tail = p.add_run(text[len(bold_prefix):])
        set_run_font(tail, size=TOKENS["body_size"], color=TOKENS["ink"], italic=italic)
    else:
        run = p.add_run(text)
        set_run_font(run, size=TOKENS["body_size"], color=TOKENS["ink"], italic=italic)
    set_keep_together(p)
    return p


def add_bullet(doc, text, current=False):
    p = doc.add_paragraph(style="Resume Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1.8)
    p.paragraph_format.line_spacing = TOKENS["body_line"]
    run = p.add_run(text)
    set_run_font(run, size=TOKENS["body_size"], color=TOKENS["ink"])
    set_cant_split(p)
    return p


def add_skill_line(doc, label, value):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(1.1)
    p.paragraph_format.line_spacing = 1.0
    add_inline(
        p,
        [
            {"text": f"{label}: ", "bold": True, "color": TOKENS["ink"]},
            {"text": value, "color": TOKENS["ink"]},
        ],
    )
    set_keep_together(p)
    return p


def add_experience(doc, company, role, dates, location, bullets):
    company_line = doc.add_paragraph(style="Heading 2")
    company_line.paragraph_format.space_before = Pt(4.5)
    company_line.paragraph_format.space_after = Pt(0.4)
    set_keep_with_next(company_line)
    add_inline(
        company_line,
        [
            {"text": company, "bold": True, "size": TOKENS["entry_size"] + 0.5, "color": TOKENS["teal"]},
            {"text": " | ", "color": TOKENS["muted"]},
            {"text": dates, "italic": True, "size": TOKENS["body_size"], "color": TOKENS["muted"]},
        ],
    )
    role_line = doc.add_paragraph(style="Resume Entry")
    role_line.paragraph_format.space_after = Pt(1.3)
    set_keep_with_next(role_line)
    add_inline(
        role_line,
        [
            {"text": role, "bold": True, "size": TOKENS["entry_size"], "color": TOKENS["ink"]},
            {"text": f" | {location}", "italic": True, "size": TOKENS["body_size"], "color": TOKENS["muted"]},
        ],
    )
    for bullet in bullets:
        add_bullet(doc, bullet)


def add_project(doc, name, descriptor, url, context, bullets, technologies, page_break_before=False):
    title = doc.add_paragraph(style="Heading 2")
    title.paragraph_format.space_before = Pt(5.1)
    title.paragraph_format.space_after = Pt(0.8)
    title.paragraph_format.page_break_before = page_break_before
    set_keep_with_next(title)
    add_inline(
        title,
        [
            {"text": f"{name} - {descriptor}", "bold": True, "size": TOKENS["entry_size"] + 0.25, "color": TOKENS["navy"]},
            {"text": " | ", "color": TOKENS["muted"]},
            {"text": url.replace("https://", ""), "url": url, "color": TOKENS["link"]},
        ],
    )
    context_line = doc.add_paragraph(style="Project Detail")
    context_line.paragraph_format.space_after = Pt(1.2)
    context_line.paragraph_format.line_spacing = 1.0
    context_run = context_line.add_run(context)
    set_run_font(context_run, size=TOKENS["body_size"], color=TOKENS["muted"], italic=True)
    set_keep_together(context_line)
    set_keep_with_next(context_line)

    for index, bullet in enumerate(bullets):
        item = add_bullet(doc, bullet)
        set_keep_with_next(item)

    tech = doc.add_paragraph(style="Project Detail")
    tech.paragraph_format.space_after = Pt(1.5)
    tech.paragraph_format.line_spacing = 1.0
    add_inline(
        tech,
        [
            {"text": "Tech Stack: ", "bold": True, "size": TOKENS["body_size"] - 0.2, "color": TOKENS["teal"]},
            {"text": technologies, "italic": True, "size": TOKENS["body_size"] - 0.2, "color": TOKENS["muted"]},
        ],
    )
    set_keep_together(tech)


def add_impact_line(doc):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_before = Pt(2.0)
    p.paragraph_format.space_after = Pt(3.0)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.line_spacing = 1.0
    add_paragraph_shading(p, TOKENS["impact_fill"])
    add_bottom_border(p, color=TOKENS["rule"], size="4", space="4")
    add_inline(
        p,
        [
            {"text": "SELECTED IMPACT  ", "bold": True, "size": TOKENS["body_size"] - 0.1, "color": TOKENS["teal"]},
            {"text": "8+ production admin panels and landing pages  |  10+ feature modules  |  approximately 80% bundle reduction", "bold": True, "size": TOKENS["body_size"], "color": TOKENS["navy"]},
        ],
    )
    set_keep_together(p)


def define_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = TOKENS["font"]
    normal._element.rPr.rFonts.set(qn("w:ascii"), TOKENS["font"])
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), TOKENS["font"])
    normal.font.size = Pt(TOKENS["body_size"])
    normal.font.color.rgb = rgb(TOKENS["ink"])
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(TOKENS["body_after"])
    normal.paragraph_format.line_spacing = TOKENS["body_line"]

    for name in ("Resume Entry", "Project Detail", "Resume Bullet"):
        if name not in styles:
            styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style = styles[name]
        style.base_style = normal
        style.font.name = TOKENS["font"]
        style._element.rPr.rFonts.set(qn("w:ascii"), TOKENS["font"])
        style._element.rPr.rFonts.set(qn("w:hAnsi"), TOKENS["font"])
        style.font.size = Pt(TOKENS["body_size"])
        style.font.color.rgb = rgb(TOKENS["ink"])

    heading_1 = styles["Heading 1"]
    heading_1.font.name = TOKENS["font"]
    heading_1._element.rPr.rFonts.set(qn("w:ascii"), TOKENS["font"])
    heading_1._element.rPr.rFonts.set(qn("w:hAnsi"), TOKENS["font"])
    heading_1.font.size = Pt(TOKENS["section_size"])
    heading_1.font.bold = True
    heading_1.font.color.rgb = rgb(TOKENS["navy"])
    heading_1.paragraph_format.space_before = Pt(7.0)
    heading_1.paragraph_format.space_after = Pt(4.2)
    heading_1.paragraph_format.keep_with_next = True

    heading_2 = styles["Heading 2"]
    heading_2.font.name = TOKENS["font"]
    heading_2._element.rPr.rFonts.set(qn("w:ascii"), TOKENS["font"])
    heading_2._element.rPr.rFonts.set(qn("w:hAnsi"), TOKENS["font"])
    heading_2.font.size = Pt(TOKENS["entry_size"])
    heading_2.font.bold = True
    heading_2.font.color.rgb = rgb(TOKENS["navy"])
    heading_2.paragraph_format.space_before = Pt(4.5)
    heading_2.paragraph_format.space_after = Pt(0.8)
    heading_2.paragraph_format.keep_with_next = True

    bullet = styles["Resume Bullet"]
    bullet.paragraph_format.left_indent = Inches(0.19)
    bullet.paragraph_format.first_line_indent = Inches(-0.13)
    bullet.paragraph_format.space_after = Pt(1.8)
    bullet.paragraph_format.line_spacing = TOKENS["body_line"]

    # Attach a real Word bullet numbering definition from the built-in style.
    builtin_bullet = styles["List Bullet"]
    num_pr = builtin_bullet._element.pPr.find(qn("w:numPr")) if builtin_bullet._element.pPr is not None else None
    if num_pr is not None:
        p_pr = bullet._element.get_or_add_pPr()
        old_num_pr = p_pr.find(qn("w:numPr"))
        if old_num_pr is not None:
            p_pr.remove(old_num_pr)
        from copy import deepcopy
        p_pr.append(deepcopy(num_pr))

    project = styles["Project Detail"]
    project.paragraph_format.left_indent = Inches(0)
    project.paragraph_format.first_line_indent = Inches(0)


def build():
    doc = Document()
    doc.core_properties.title = "Hamid Yeganeh - Front-End Developer Resume"
    doc.core_properties.subject = "ATS-optimized resume for Front-End Developer and React Developer roles"
    doc.core_properties.author = "Hamid Yeganeh"
    doc.core_properties.keywords = "Front-End Developer, React Developer, React, Next.js, TypeScript"
    define_styles(doc)

    section = doc.sections[0]
    section.page_width = Inches(TOKENS["page_width"])
    section.page_height = Inches(TOKENS["page_height"])
    section.top_margin = Inches(TOKENS["margin"])
    section.bottom_margin = Inches(TOKENS["margin"])
    section.left_margin = Inches(TOKENS["margin"])
    section.right_margin = Inches(TOKENS["margin"])
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    # ATS-safe memo_masthead adaptation: all identity/contact content is in the
    # document body, never in a header, footer, text box, column, or table.
    name = doc.add_paragraph()
    name.paragraph_format.space_before = Pt(0)
    name.paragraph_format.space_after = Pt(0.5)
    set_keep_with_next(name)
    run = name.add_run("Hamid Yeganeh")
    set_run_font(run, size=TOKENS["title_size"], color=TOKENS["ink"], bold=True)

    role = doc.add_paragraph()
    role.paragraph_format.space_after = Pt(1.5)
    set_keep_with_next(role)
    run = role.add_run("Front-End Developer | React & Next.js Developer")
    set_run_font(run, size=TOKENS["role_size"], color=TOKENS["navy"], bold=True)

    contact = doc.add_paragraph()
    contact.paragraph_format.space_after = Pt(4.5)
    contact.paragraph_format.line_spacing = 1.0
    add_inline(
        contact,
        [
            {"text": "+98 938 372 9627 | ", "color": TOKENS["ink"]},
            {"text": "hamidyeganeh82@gmail.com", "url": "mailto:hamidyeganeh82@gmail.com"},
            {"text": " | ", "color": TOKENS["muted"]},
            {"text": "github.com/hamidYeganeh", "url": "https://github.com/hamidYeganeh"},
            {"text": " | ", "color": TOKENS["muted"]},
            {"text": "linkedin.com/in/hamidYeganeh", "url": "https://www.linkedin.com/in/hamidYeganeh"},
            {"text": " | Tehran, Iran", "color": TOKENS["ink"]},
        ],
    )
    add_bottom_border(contact, color=TOKENS["teal"], size="7", space="5")

    add_section_heading(doc, "Professional Summary")
    add_body(
        doc,
        "Front-End Developer with 4+ years of experience delivering production web applications, operational dashboards, admin platforms, and public websites with React, Next.js, and TypeScript. Develops frontend features end-to-end across architecture, responsive UI, API integration, state management, access control, testing, production monitoring, and deployment. Built and maintained 8+ production admin panels and landing pages and contributed to large multi-module platforms with real-time and multilingual experiences. Reduced one platform's client bundle from 700 KB to 140 KB (approximately 80%) through code splitting, lazy loading, targeted refactoring, and bundle analysis."
    )
    add_impact_line(doc)

    add_section_heading(doc, "Technical Skills")
    add_skill_line(doc, "Core", "HTML5, CSS3, JavaScript (ES6+), TypeScript")
    add_skill_line(doc, "Frameworks/Libraries", "React, Next.js (App Router, Pages Router, SSR, SSG, SEO), Redux, Redux Toolkit, Zustand, RTK Query, React Query, React Hook Form, Zod, next-intl, Socket.IO")
    add_skill_line(doc, "Styling", "Tailwind CSS, Sass/SCSS, Styled Components, shadcn/ui, Radix UI, Material UI, Ant Design")
    add_skill_line(doc, "Tools", "Git, GitHub, GitLab, Storybook, Turborepo, NX, Figma, Docker, CI/CD, Sentry, Postman, Swagger")
    add_skill_line(doc, "Testing/Performance", "Jest, React Testing Library, Lighthouse, code splitting, lazy loading, memoization, bundle analysis")
    add_skill_line(doc, "Frontend Practices", "Responsive Web Design, REST API integration, RBAC, JWT/OAuth, PWA, Service Workers, feature-based architecture, monorepo, component-driven development, i18n/localization")

    add_section_heading(doc, "Professional Experience")
    add_experience(
        doc,
        "Bareach",
        "Front-End Developer",
        "Mar 2026 - Present",
        "Tehran, Iran",
        [
            "Develop and maintain a transportation operations dashboard with Next.js App Router and TypeScript, translating complex workflows into a feature-based frontend with 10+ independently maintainable modules.",
            "Design multi-step dynamic forms with React Hook Form and Zod, using shared validation patterns to keep field behavior consistent across critical operational flows.",
            "Manage server and client state with React Query and Zustand, supporting reliable API synchronization and predictable UI updates across dashboard modules.",
            "Strengthen production reliability and access control through RBAC guards, Sentry monitoring, next-intl localization, and base unit tests with Jest and React Testing Library.",
        ],
    )
    add_experience(
        doc,
        "Arshco",
        "Front-End Developer (promoted from Intern in 6 months)",
        "Dec 2021 - Mar 2026",
        "Tehran, Iran",
        [
            "Promoted from Front-End Intern to Front-End Developer after six months and continued contributing across a four-year tenure.",
            "Built and maintained 8+ production admin panels and landing pages while working directly with backend engineers and product owners to deliver scoped features.",
            "Developed reusable UI component libraries that standardized interaction and visual patterns across products and simplified implementation of recurring interfaces.",
            "Integrated REST APIs and implemented authentication and authorization workflows with JWT and OAuth across multiple production applications.",
        ],
    )

    add_section_heading(doc, "Education")
    edu = doc.add_paragraph(style="Resume Entry")
    edu.paragraph_format.space_before = Pt(2)
    edu.paragraph_format.space_after = Pt(0)
    add_inline(
        edu,
        [
            {"text": "Shamsipour Technical and Vocational College", "bold": True, "size": TOKENS["entry_size"], "color": TOKENS["navy"]},
            {"text": " | Bachelor of Science in Computer Engineering (in progress)", "color": TOKENS["ink"]},
        ],
    )

    add_section_heading(doc, "Selected Projects", page_break_before=True)
    add_project(
        doc,
        "BizMLM",
        "Network Marketing & Commerce Platform",
        "https://bizmlm.ir",
        "Large multi-area platform spanning an admin panel, user dashboard, academy, e-commerce shop, articles, and real-time product experiences.",
        [
            "Contributed across 6+ frontend modules inside a Turborepo monorepo; migrated legacy pages to a modern Next.js architecture, implemented next-intl localization, and documented reusable UI components in Storybook.",
            "Reduced the client bundle from 700 KB to 140 KB (approximately 80%) through code splitting, lazy loading, targeted refactoring, and deep bundle analysis, improving initial delivery efficiency.",
        ],
        "Next.js Pages Router, TypeScript, Redux, Styled Components, Turborepo, Socket.IO, Storybook, ECharts, GSAP, PWA.",
    )
    add_project(
        doc,
        "Eqipo",
        "Real Estate Crowdfunding Platform",
        "https://eqipo.ir",
        "Real-estate crowdfunding product combining public acquisition pages, investor-facing workflows, an authenticated dashboard, and editorial content.",
        [
            "Built responsive landing, dashboard, blog, and content experiences with a consistent component system across public and authenticated areas.",
            "Implemented RBAC for permission-based workflows, rich investment-data tables with sorting, filtering, and pagination, plus PWA service workers and offline support.",
        ],
        "Next.js App Router, TypeScript, Tailwind CSS, Redux Toolkit, Radix UI, React Table, Framer Motion, PWA.",
    )
    add_project(
        doc,
        "Bareach",
        "Transportation Operations Dashboard",
        "https://bareach.com",
        "Operational transportation dashboard with complex workflows, dynamic forms, permission-sensitive modules, synchronized API data, and production observability requirements.",
        [
            "Architected 10+ feature modules and implemented dynamic multi-step forms with shared Zod validation, while coordinating server and client state through React Query and Zustand.",
            "Added RBAC guards, Sentry monitoring, next-intl localization, Framer Motion interactions, and critical component tests to support a secure and maintainable production experience.",
        ],
        "Next.js App Router, TypeScript, Tailwind CSS, shadcn/ui, React Hook Form, Zod, React Query, Zustand, Sentry, Jest, React Testing Library.",
    )
    add_project(
        doc,
        "Greenkeeper",
        "Environmental Certificate Platform",
        "https://greenkeeper.eco",
        "International environmental certificate platform with distinct user, administrator, organization, and public landing experiences connected to a microservice backend.",
        [
            "Developed four product areas and a reusable Tailwind CSS and shadcn/ui component foundation, maintaining consistent behavior across different user roles.",
            "Integrated multiple backend services with RTK Query and delivered German and English localization for international product and content flows.",
        ],
        "Next.js Pages Router, TypeScript, Tailwind CSS, shadcn/ui, Redux, RTK Query, Framer Motion, microservice architecture.",
    )
    add_project(
        doc,
        "Royan Cancer",
        "Medical Research Institute Website",
        "https://royancancer.org/",
        "Public-facing medical research website focused on discoverability, readable content presentation, responsive behavior, and maintainable article publishing.",
        [
            "Built responsive landing and dynamic article listing/detail pages with server-side rendering, structured content handling, and maintainable page composition.",
            "Applied accessible layout patterns and SSR to support readable experiences across devices, stronger search visibility, and faster initial content rendering.",
        ],
        "Next.js App Router, Tailwind CSS, server-side rendering, dynamic content.",
    )

    # Prevent Word from adding compatibility spacing drift.
    settings = doc.settings._element
    compat = settings.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        settings.append(compat)
    setting = OxmlElement("w:doNotUseHTMLParagraphAutoSpacing")
    compat.append(setting)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
